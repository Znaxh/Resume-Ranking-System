import os
import tempfile
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document
import logging
import time

import app_config

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handle file upload and text extraction"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.allowed_extensions = set(app_config.ALLOWED_EXTENSIONS)
        self.max_file_size = app_config.MAX_FILE_SIZE_MB * 1024 * 1024
    
    def validate_file(self, file) -> Dict[str, Any]:
        """Validate uploaded file"""
        if not file or not file.filename:
            return {'valid': False, 'error': 'No file provided'}
        
        # Check file extension
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_extension not in self.allowed_extensions:
            return {
                'valid': False, 
                'error': f'File type .{file_extension} not supported. Allowed: {", ".join(self.allowed_extensions)}'
            }
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > self.max_file_size:
            return {
                'valid': False,
                'error': f'File too large. Maximum size: {self.max_file_size // (1024*1024)}MB'
            }
        
        return {
            'valid': True,
            'filename': file.filename,
            'size': file_size,
            'extension': file_extension
        }
    
    def extract_text_from_pdf(self, file) -> str:
        """Extract text from PDF file - Windows compatible version"""
        temp_file_path = None
        try:
            # Create temporary file with delete=False for Windows compatibility
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file_path = temp_file.name
                
                # Reset file pointer
                file.seek(0)
                
                # Write file content to temp file
                temp_file.write(file.read())
                temp_file.flush()
            
            # Small delay to ensure file is written completely
            time.sleep(0.1)
            
            # Extract text from the temporary file
            text = ""
            with open(temp_file_path, 'rb') as pdf_file:
                try:
                    pdf_reader = PdfReader(pdf_file)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                            continue
                    
                except Exception as e:
                    logger.error(f"PDF reader error: {e}")
                    # Fallback: try alternative method
                    text = self._extract_pdf_fallback(temp_file_path)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
        
        finally:
            # Clean up temp file
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    time.sleep(0.1)  # Small delay for Windows
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file {temp_file_path}: {e}")
    
    def _extract_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF extraction method"""
        try:
            import fitz  # PyMuPDF - alternative PDF reader
            
            doc = fitz.open(file_path)
            text = ""
            
            for page in doc:
                text += page.get_text() + "\n"
            
            doc.close()
            return text
            
        except ImportError:
            logger.warning("PyMuPDF not available for fallback PDF extraction")
            return "Unable to extract PDF text"
        except Exception as e:
            logger.error(f"Fallback PDF extraction failed: {e}")
            return "Unable to extract PDF text"
    
    def extract_text_from_docx(self, file) -> str:
        """Extract text from DOCX file - Windows compatible version"""
        temp_file_path = None
        try:
            # Create temporary file with delete=False for Windows compatibility
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file_path = temp_file.name
                
                # Reset file pointer
                file.seek(0)
                
                # Write file content to temp file
                temp_file.write(file.read())
                temp_file.flush()
            
            # Small delay to ensure file is written completely
            time.sleep(0.1)
            
            # Extract text from the temporary file
            doc = Document(temp_file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
        
        finally:
            # Clean up temp file
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    time.sleep(0.1)  # Small delay for Windows
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file {temp_file_path}: {e}")
    
    def _extract_text_from_legacy_doc(self, file) -> str:
        """Extract text from legacy .doc files using subprocess fallback"""
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file_path = temp_file.name
                file.seek(0)
                temp_file.write(file.read())
                temp_file.flush()

            import subprocess
            try:
                result = subprocess.run(
                    ['antiword', temp_file_path],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except FileNotFoundError:
                pass

            try:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir',
                     os.path.dirname(temp_file_path), temp_file_path],
                    capture_output=True, text=True, timeout=60
                )
                txt_path = temp_file_path.rsplit('.', 1)[0] + '.txt'
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()
                    os.unlink(txt_path)
                    if text.strip():
                        return text.strip()
            except FileNotFoundError:
                pass

            logger.warning("Neither antiword nor libreoffice available for .doc conversion, attempting python-docx fallback")
            return self.extract_text_from_docx(file)

        except Exception as e:
            logger.error(f"Legacy .doc extraction failed: {e}")
            raise Exception(
                "Failed to extract text from .doc file. "
                "For best results, convert to .docx or .pdf format."
            )
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass

    def process_files(self, files: List) -> List[Dict[str, Any]]:
        """Process multiple uploaded files"""
        processed_files = []
        
        for i, file in enumerate(files):
            try:
                # Validate file
                validation_result = self.validate_file(file)
                
                if not validation_result['valid']:
                    processed_files.append({
                        'index': i,
                        'filename': file.filename if file else f'file_{i}',
                        'success': False,
                        'error': validation_result['error'],
                        'text': ""
                    })
                    continue
                
                # Reset file pointer before processing
                file.seek(0)
                
                # Extract text based on file extension
                file_extension = validation_result['extension']
                
                logger.info(f"Processing file {i+1}: {validation_result['filename']} ({file_extension})")
                
                if file_extension == 'pdf':
                    text = self.extract_text_from_pdf(file)
                elif file_extension == 'docx':
                    text = self.extract_text_from_docx(file)
                elif file_extension == 'doc':
                    text = self._extract_text_from_legacy_doc(file)
                elif file_extension == 'txt':
                    file.seek(0)
                    text = file.read().decode('utf-8', errors='replace')
                else:
                    raise Exception(f"Unsupported file type: {file_extension}")
                
                # Validate extracted text
                if not text or len(text.strip()) < 50:
                    processed_files.append({
                        'index': i,
                        'filename': validation_result['filename'],
                        'success': False,
                        'error': 'File appears to be empty or contains insufficient text',
                        'text': ""
                    })
                    continue
                
                processed_files.append({
                    'index': i,
                    'filename': validation_result['filename'],
                    'success': True,
                    'text': text,
                    'size': validation_result['size'],
                    'word_count': len(text.split()),
                    'char_count': len(text)
                })
                
                logger.info(f"Successfully processed file: {validation_result['filename']} ({len(text)} chars)")
                
            except Exception as e:
                logger.error(f"Error processing file {i}: {e}")
                processed_files.append({
                    'index': i,
                    'filename': file.filename if file and file.filename else f'file_{i}',
                    'success': False,
                    'error': str(e),
                    'text': ""
                })
        
        return processed_files
